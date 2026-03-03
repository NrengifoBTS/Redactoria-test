#redactoria/src/scraping/service.py
import re
import json
import requests
import spacy
from bs4 import BeautifulSoup, Tag
from collections import Counter
from typing import Generator, List, Dict, Any, Optional, Union,Set
from datetime import datetime, timezone
from . import models 
from pydantic import BaseModel
import urllib3
from fastapi import HTTPException
from .models import PeticionGeneracionContenido
from fastapi.responses import StreamingResponse, FileResponse
from src.entities.scraping import Scraping
from uuid import UUID
import logging
from sqlalchemy.orm import Session

import cloudscraper
from fake_useragent import UserAgent
import trafilatura

from docx import Document
from docx.shared import Inches
from io import BytesIO

# Importación necesaria para la persistencia
from src.entities.scraping import Scraping
from src.entities.blog import Blog

# =======================================================================
# GENERACION DE DOCUMENTO WORD 
# =======================================================================

class DocumentService:
    """
    Servicio para generar documentos Word (.docx) a partir de la estructura del blog.
    """

    # Atributo de clase. Se accede vía self.MAPEO_NIVELES o DocumentService.MAPEO_NIVELES
    MAPEO_NIVELES = {
        'h1': 1, # Nivel 1 en Word (equivale al estilo 'Heading 1', el Título 1)
        'h2': 2, # Nivel 2 en Word (equivale al estilo 'Heading 2', el Título 2)
        'h3': 3,
        'h4': 4,
        'h5': 5,
        'h6': 6,
    }

    def procesar_seccion_recursivamente(
        self,
        item_seccion: Dict[str, Any], 
        documento_word: Document
    ):
        """
        Función recursiva: Procesa un elemento (título y contenido) y luego sus hijos.
        """
        
        nivel_str = item_seccion.get('level', 'h2').lower()
        titulo_seccion = item_seccion.get('text') 
        
        # Conversión del nivel (accediendo al mapa)
        # 💡 Si el nivel no existe ('h7', etc.), se asigna un nivel por defecto (ej. 2 para H2).
        nivel_encabezado = self.MAPEO_NIVELES.get(nivel_str, 2) 

        # 1. Añadir título
        if titulo_seccion:
            # Ahora, h1 usará level=1 (Heading 1) y h2 usará level=2 (Heading 2)
            documento_word.add_heading(titulo_seccion, level=nivel_encabezado)
        
        # 2. Añadir contenido (el cuerpo del texto)
        contenido_cuerpo = item_seccion.get('content')
        if contenido_cuerpo and contenido_cuerpo.strip(): 
            documento_word.add_paragraph(contenido_cuerpo)
            
        # 3. Procesar hijos (recursividad)
        if item_seccion.get('children'):
            for hijo in item_seccion['children']:
                # Llamada recursiva
                self.procesar_seccion_recursivamente(hijo, documento_word)


    def generar_documento_word(
        self,
        blog_id: UUID, 
        datos_estructura: List[Any], 
        db: Session,
    ) -> StreamingResponse:
        
        # 1. Verificar existencia
        blog = db.query(Blog).filter(Blog.id == blog_id).first()
        if not blog:
            raise HTTPException(status_code=404, detail="Blog no encontrado.")
            
        # 2. Creación del Documento Word
        documento = Document()
        
        # 3. Mapeo de la estructura: Procesar la lista completa enviada
        for item_principal in datos_estructura:
            try:
                # Conversión de Pydantic a diccionario
                # Si item_principal es BlogSectionData, .model_dump() es correcto.
                item_principal_dict = item_principal.model_dump() 
                
                # LA CLAVE: Aquí se llama a la función recursiva, 
                # la cual utiliza el MAPEO_NIVELES corregido.
                self.procesar_seccion_recursivamente(item_principal_dict, documento)
                
            except Exception as e:
                # Usar logging.error o simplemente raise el error si es crítico.
                logging.error(f"Error al procesar item de estructura: {e}. Item: {item_principal}")
                continue

        # 4. Guardar y Retornar
        flujo_archivo = BytesIO()
        documento.save(flujo_archivo)
        flujo_archivo.seek(0)

        # -----------------------------------------------------------------
        
        return StreamingResponse(
            flujo_archivo,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={} 
        )



# =======================================================================
# FUNCIONES DE PERSISTENCIA DE LA TABLA SCRAPING Y BLOG
# =======================================================================
def actualizar_o_crear_resultado_scraping(
    db: Session, 
    blog_id: UUID, 
    datos_a_guardar: Dict[str, Any]
) -> Scraping:
    """
    Crea o actualiza de forma genérica una entidad ScrapingResult 
    para un blog específico con los datos proporcionados.
    """
    
    # 1. Buscar si ya existe un resultado para este blog_id
    resultado_scraping = db.query(Scraping).filter(
        Scraping.blog_id == blog_id
    ).first()

    if resultado_scraping:
        # --- Lógica de ACTUALIZACIÓN ---
        
        # CRÍTICO: Asignar el campo consolidated_content si está presente
        if 'consolidated_content' in datos_a_guardar:
            resultado_scraping.consolidated_content = datos_a_guardar['consolidated_content']
            
        # Asignar otros campos genéricos si están en el diccionario
        for key, value in datos_a_guardar.items():
            if hasattr(resultado_scraping, key):
                setattr(resultado_scraping, key, value)
                
    else:
        # --- Lógica de CREACIÓN ---
        # Asegúrate de incluir blog_id y el contenido consolidado en el objeto de creación
        datos_iniciales = {
            "blog_id": blog_id,
            "consolidated_content": datos_a_guardar.get('consolidated_content', None),
            # Incluir otros campos requeridos por la entidad Scraping
            **datos_a_guardar 
        }
        
        resultado_scraping = Scraping(**datos_iniciales)
        db.add(resultado_scraping)
    
    db.commit()
    db.refresh(resultado_scraping)
    return resultado_scraping

def actualizar_estructura_blog(
    db: Session, 
    blog_id: UUID, 
    estructura_data: Dict[str, Any],
    estimated_word_count: Optional[int] = None # <-- NUEVO PARÁMETRO
) -> Blog:
    
    blog = db.query(Blog).filter(Blog.id == blog_id).first()
    # ... (lógica de manejo de error) 

    # 1. Actualizar estructura (lógica existente)
    blog.estructura_blog_json = estructura_data
    
    # 2. LÓGICA CRÍTICA: Guardar el conteo en la tabla Blog
    if estimated_word_count is not None:
        # El nombre del campo en la entidad Blog ahora debe ser 'estimated_word_count'
        blog.estimated_word_count = estimated_word_count 
    
    db.commit()
    db.refresh(blog)
    
    return blog



urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
# --- 1. CLASE AIService: Control y Generación de IA ---

class AIService:
    """
    Gestiona todas las interacciones con el modelo de lenguaje (LLM),
    incluyendo resúmenes de bloques y análisis final.
    """

    #MODEL_URL = "http://192.168.1.36:1234/v1/chat/completions" #<-- Compu Alda
    MODEL_URL = "http://host.docker.internal:1234/v1/chat/completions" 
    MODEL_NAME = "redactoria-v3-gold"
    DEFAULT_SYSTEM_MESSAGE = (
        "Eres un Redactor SEO, Copywriter y Editor Web de ÉLITE. "
        "INSTRUCCIÓN CRÍTICA: NO USEA NINGUNA TIPOGRAFIA ESPECIAL ,NEGRITAS, CURSIVAS , SUBRRAYADO EN NINGUNA GENERACION"
        
    )

    def __init__(self):
        pass

    # LLAMADA A LM STUDIO
    def _llm_generate(self, 
                      prompt: str, 
                      system_message: str = DEFAULT_SYSTEM_MESSAGE, 
                      temperature: float = 0.4, 
                      max_tokens: Optional[int] = None 
                      ) -> str:
        
        data = {
            "model": self.MODEL_NAME,
            "messages": [
                {"role": "system", "content": system_message},
                {"role": "user", "content": prompt}
            ],
            "temperature":temperature ,
            "stream": False
        }
        
        UNIVERSAL_MAX_TOKENS = 10000 

        data["max_tokens"] = UNIVERSAL_MAX_TOKENS
        try:
            response = requests.post(self.MODEL_URL, headers={"Content-Type": "application/json"}, json=data)
            response.raise_for_status() 
            return response.json()["choices"][0]["message"]["content"].strip()
        
        except requests.exceptions.RequestException as e:
            error_message = f"[FALLO LLM - {response.status_code if 'response' in locals() else 'Red'}: {type(e).__name__} - {str(e)}]"
            return error_message 
        except Exception as e:
            return f"[Error interno inesperado: {e}]"

    

    # PARSEO DE FORMATO JSON DE LA RESPUESTA IA 
    def limpieza_extraccion_json(self, json_string: str) -> Dict[str, Any]:
    
        # ------------------------------------------------------------------
        # FUNCIÓN ANIDADA: Saneamiento agresivo de caracteres que rompen JSON
        # ------------------------------------------------------------------
        def sanitize_content_for_json(data: Dict[str, Any]) -> Dict[str, Any]:
            """
            Itera sobre el diccionario y aplica un saneamiento doble a las cadenas de texto:
            1. Escapa backslashes ( \ -> \\\\ ).
            2. Elimina caracteres de control ilegales en JSON.
            """
            sanitized_data = {}
            
            for key, value in data.items():
                if isinstance(value, str):
                    safe_value = value.replace('\\', '\\\\')

                    control_char_pattern = r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]' 
                    safe_value = re.sub(control_char_pattern, '', safe_value)
                    
                    sanitized_data[key] = safe_value
                else:
                    sanitized_data[key] = value
            return sanitized_data
        # ------------------------------------------------------------------


        # 1. Limpieza inicial para remover etiquetas Markdown (```json)
        clean_json = re.sub(r'```json\s*|```', '', json_string, flags=re.IGNORECASE).strip()
    
        try:
            # 2. Intento de parseo estándar
            parsed_data = json.loads(clean_json)
            
            # 3. Saneamiento agresivo y retorno del resultado válido
            return sanitize_content_for_json(parsed_data)
            
        except json.JSONDecodeError as e:
            # Reemplazar comillas simples por dobles 
            temp_clean_json = clean_json.replace("'", '"')
            temp_clean_json = re.sub(r'"\s*\n\s*"', '",\n"', temp_clean_json)
            temp_clean_json = re.sub(r'\}\s*\"', '},\n\"', temp_clean_json)
            temp_clean_json = re.sub(r',\s*\}', '}', temp_clean_json)
            temp_clean_json = re.sub(r',\s*\]', ']', temp_clean_json)
            temp_clean_json = re.sub(r'(\s*[\}\]])\s*(\s*[\{\[])', r'\1,\2', temp_clean_json)
            start_index = temp_clean_json.find('{')
            end_index = temp_clean_json.rfind('}')
            
            if start_index != -1 and end_index != -1:
                try:
                    # 1. Parseo forzado sobre la cadena saneada y delimitada
                    forced_data = json.loads(temp_clean_json[start_index:end_index + 1])
                    
                    # 2. Saneamiento de Backslashes y retorno
                    return sanitize_content_for_json(forced_data)
                    
                except json.JSONDecodeError as nested_e:
                    # 3. Si falla incluso el parseo forzado con corrección de sintaxis, es irrecuperable.
                    raise ValueError(
                        f"Fallo de parseo JSON irrecuperable. El LLM no devolvió JSON válido. "
                        f"Error final: {nested_e}"
                    ) from nested_e
            
            # Si no se encontró la estructura { ... } después del fallo inicial.
            raise ValueError(
                f"El LLM devolvió contenido sin estructura JSON ({{...}}). "
                f"Inicio del output: {clean_json[:200]}..."
            ) from e


    #SE ENCARGA DE FORMATEAR LA INFORMACION EXTRAIDA Y UTILIZARLA POR EL LLM 
    def _build_media_text(self, media_info: List[Dict[str, Any]]) -> str:
        if not media_info:
            return "No se encontró contenido multimedia relevante para este bloque."

        media_list = [
            f"- Tipo: {item.get('type', 'Imagen')}, Descripción: {item.get('alt', 'No proporcionada')}"
            for item in media_info
        ]

        media_text = "\n".join(media_list)
        return f"\n--- REFERENCIA DE CONTENIDO MULTIMEDIA (USAR SOLO COMO CONTEXTO) ---\n{media_text}\n---\n"

    def estimar_longitud(self, db: Session, blog_id: UUID, query: str) -> tuple[int, str]:
        """
        Orquesta la recuperación de longitudes de la competencia y llama a la IA 
        para juzgar la longitud óptima del artículo.
        Retorna: (longitud_estimada_int, longitudes_competencia_str)
        """
        # Se necesita la entidad ScrapeResult para obtener los conteos de palabras por URL
        from src.entities.scraping import ScrapeResult 
        
        longitudes_competencia_str = 'N/A'
        judged_word_count = 1500 # Valor de fallback seguro
        
        try:
            # 1. Recuperar longitudes de la competencia (asumiendo que ScrapeResult guarda el conteo)
            scrape_results_list = db.query(ScrapeResult).filter(ScrapeResult.blog_id == blog_id).all()
            
            if scrape_results_list:
                longitudes = [
                    str(result.word_count) 
                    for result in scrape_results_list 
                    if result.word_count is not None and result.word_count > 0
                ]
                if longitudes:
                    longitudes_competencia_str = ", ".join(longitudes)

            # 2. Juicio de la IA solo si hay datos válidos de competencia
            if longitudes_competencia_str != 'N/A':
                
                system_msg_juicio = "Eres un Analista SEO experto. Tu única tarea es generar una estimación de la longitud de un artículo para posicionamiento en Google."
                
                prompt_juicio = f"""
                    Como analista experto, tu tarea es determinar la longitud de palabras óptima para un nuevo 
                    artículo SEO sobre el tema: '{query}'.

                    Las longitudes de los artículos de la competencia (incluyendo contenido sucio) son: {longitudes_competencia_str}.

                    Usando tu juicio para descartar posibles outliers o contenido irrelevante, determina una 
                    longitud estimada óptima.

                    **CRÍTICO: Devuelve ÚNICAMENTE el número entero de palabras, sin comas, puntos, ni texto adicional. 
                    Ejemplo de salida correcta: 1850**
                """
                
                result_str = self._llm_generate(
                    prompt_juicio, 
                    system_msg_juicio, 
                    temperature=0.1, 
                    max_tokens=10 
                )

                # 3. Limpieza robusta del resultado
                cleaned_number_str = re.sub(r'[^\d]', '', result_str).strip()
                
                if cleaned_number_str:
                    judged_word_count = int(cleaned_number_str)

        except Exception as e:
            print(f"ERROR en estimar_longitud para blog {blog_id}: {e}") 

        print(f"Longitud estimada (Juicio de IA): {judged_word_count}. Longitudes de competencia: {longitudes_competencia_str}")
        return judged_word_count, longitudes_competencia_str


    # --- ANALISIS DE BLOQUES CON IA ---
    def analizar_bloque_contenido(self, chunk: str, media_info: List[Dict[str, str]], query: str, heading: str, contexto_previo: str = "") -> str:
        MAX_CHARS_PER_LLM_CALL = 6000
        media_text = self._build_media_text(media_info)

        # 1. OPTIMIZACIÓN DE MEMORIA: Solo extraemos nombres de entidades/lugares específicos ya mencionados
        # para que la IA no repita el MISMO restaurante, pero sí analice la fuente.
        subtemas_previos = re.findall(r"- Subtema: (.*)", contexto_previo)
        lista_temas_evitar = ", ".join(set(subtemas_previos[-30:])) 
        
        memoria_ia = ""
        if lista_temas_evitar:
            # Cambiamos el enfoque: No le decimos que evite el tema, sino que evite REPETIR estas entidades exactas
            memoria_ia = f"\nENTIDADES YA REGISTRADAS (No repetir si son exactamente las mismas): {lista_temas_evitar}\n"

        def _generate_prompt(sub_chunk: str) -> str:
            return f"""Analiza este contenido para un artículo sobre '{query}':
            ---
            {sub_chunk}
            {media_text}
            ---
            {memoria_ia}

            INSTRUCCIONES DE EXTRACCIÓN:
            1. Misión: Extraer TODOS los datos, nombres propios, precios, ubicaciones y descripciones técnicas.
            2. CRITERIO DE NO REPETICIÓN: Solo ignora un dato si es EXACTAMENTE el mismo nombre y detalle que ya aparece en 'ENTIDADES YA REGISTRADAS'. 
            3. Si la fuente menciona un lugar nuevo o un detalle diferente de un tema ya cubierto, DEBES INCLUIRLO.
            4. FORMATO: Devuelve la información organizada por 'Subtema' (el nombre del lugar o concepto) y 'Hecho' (los datos).
            5. PROHIBIDO: No uses comillas, negritas, ni introducciones. Empieza directo con '- Subtema:'.

            FORMATO DE SALIDA:
            - Subtema: (Nombre)
            - Hecho: (Datos densos)
            """

        system_msg = "Eres un Documentalista de Datos. Tu objetivo es extraer el 100% de la información única de la fuente, evitando solo la duplicidad exacta de registros anteriores."
        
        paragraphs = chunk.split('\n\n')
        current_sub = ""
        bloques_respuesta = []

        for p in paragraphs:
            if len(current_sub) + len(p) > MAX_CHARS_PER_LLM_CALL:
                if current_sub.strip():
                    res = self._llm_generate(_generate_prompt(current_sub), system_msg, 0.2)
                    if res and '[FALLO' not in res: 
                        bloques_respuesta.append(res)
                current_sub = p + "\n\n"
            else:
                current_sub += p + "\n\n"
        
        if current_sub.strip():
            res = self._llm_generate(_generate_prompt(current_sub), system_msg, 0.2)
            if res and '[FALLO' not in res: 
                bloques_respuesta.append(res)

        # 2. PROCESAMIENTO FINAL: Más flexible
        final_output = []
        for bloque in bloques_respuesta:
            # Capturamos líneas que empiecen con "- " pero somos menos estrictos con la longitud
            # para no perder datos cortos pero valiosos.
            lines = [l.strip() for l in bloque.split('\n') if l.strip().startswith('- ')]
            final_output.extend(lines)

        return "\n".join(final_output).strip()
        
    def generar_estructura_seo_final(self, query: str, title_base: str, categoria: str, idioma: str, tecnica: str, acento: str, tono: str, 
                                 consolidated_text: str, estimated_word_count: int) -> Dict[str, Any]:
        """
        Genera la estructura SEO final procesando el 100% del contenido consolidado 
        en fragmentos para evitar errores de contexto (Error 400).
        """

        if not consolidated_text:
             return {
                "structure_markdown": "[ERROR: No hay texto consolidado]",
                "estimated_word_count": 0
            }

        # --- PASO 1: EXTRACCIÓN DETALLADA POR BLOQUES (SIN PÉRDIDA) ---
        # Dividimos el texto en bloques de 10,000 caracteres para procesar TODO
        chunk_size = 6000
        text_chunks = [consolidated_text[i:i+chunk_size] for i in range(0, len(consolidated_text), chunk_size)]
        
        puntos_clave_extraidos = []
        print(f"Analizando {len(text_chunks)} bloques de investigación para la estructura...")
        

        for idx, chunk in enumerate(text_chunks):
            # Prompt de extracción intermedio (no es el final, solo para recolectar datos)
            prompt_extraccion = f"""
            Analiza este fragmento ({idx+1}/{len(text_chunks)}) sobre '{title_base}':
            ---
            {chunk}
            ---
            TAREA: Extrae la información de forma ultra-comprimida siguiendo estrictamente este formato:
            - ENTIDADES: [Solo nombres propios y keywords clave separados por comas]
            - DATOS_TECNICOS: [Cifras, precios, horarios o hechos breves]
            - H2 Y H3 SUGERIDOS: [Solo títulos de secciones sugeridas para la estructura, sin descripciones]
            
            REGLA: No uses tablas, no saludes, no repitas información que ya sea obvia. Solo datos crudos.
            """
            res = self._llm_generate(prompt_extraccion, "Eres un analista de datos SEO.", temperature=0.2)
            puntos_clave_extraidos.append(res)

        # Unimos todos los puntos extraídos (esto ya es mucho más ligero que el texto bruto)
        contexto_completo_para_ia = "\n".join(puntos_clave_extraidos)
        print("Extracción de puntos clave completada." + contexto_completo_para_ia)
    
        # --- PASO 2: PERFIL EDITORIAL ACTIVO ---
        EDITORIAL_PROFILES = {
            "viajemos": """
                PERFIL: VIAJEMOS (Arquitectura de Experiencias y Viajes)
                    - PROHIBIDO crear H2 para un solo lugar si existen otros del mismo tipo.
                    - Reducción Técnica: Minimiza el uso de jerga técnica. Enfócate en la experiencia del usuario y en consejos prácticos.
                    - Soluciona la intención de búsqueda con H2 claros y específicos sin extenderse a otros temas no relacionados, solamente respondemos la intencion de busqueda.
                    - La voz del articulo no puede guiarse de a donde va o a donde viene el usuario , evita mensionar otros paises o ciudades que no esten en la intencion de busqueda
                    - En caso de ser itinerarios, cada H2 debe ser un día del itinerario y los H3 actividades dentro de ese día.
                    TONO Y ESTILO: Voz de guía experto, entusiasta y sofisticado. Títulos limpios, directos, con verbos de acción y sin signos de puntuación innecesarios.
            """,

            "arriendo": """
                PERFIL: ARRIENDO (Arquitectura Práctica y Comercial)
                - Foco: Intención de búsqueda transaccional y resolutiva.
                - Mandato de Estructura: Estructura lineal y lógica (Requisitos > Proceso > Costos > Consejos).
                - Prioridad: Claridad absoluta en los pasos a seguir. Cada H2 debe ser una etapa del proceso.
                - Reducción Técnica: No uses storytelling. Los títulos deben ser descriptivos y directos (ej. 'Documentos necesarios para arrendar').
                - Estilo de Títulos: Informativos y sobrios. Usa (LISTA) frecuentemente para procesos.
            """,

            "guia_legal": """
                PERFIL: GUÍA LEGAL (Arquitectura de Autoridad y Normativa)
                - Foco: Intención de búsqueda de consulta y seguridad jurídica.
                - Mandato de Estructura: Jerarquía basada en conceptos legales y su aplicación práctica.
                - Prioridad: Definición clara del marco legal en los primeros encabezados.
                - Reducción Técnica: La precisión es clave, pero la estructura debe ser fácil de navegar. Separa la 'Teoría' de la 'Acción' (ej. 'Qué dice la ley' vs 'Qué debes hacer tú').
                - Estilo de Títulos: Formales, técnicos pero comprensibles. Usa (TABLA) para comparar leyes o plazos.
            """
        }
        
        cat_key = categoria.lower().replace(" ", "_")
        perfil_activo = EDITORIAL_PROFILES.get(cat_key, EDITORIAL_PROFILES.get("viajemos"))

        # -- SYSTEM MESSAGE: contexto base del modelo --
        system_message = f"""
            "Eres un Arquitecto de estructuras para articulos. 

            APLICARÁS ESTE PERFIL EDITORIAL DE FORMA OBLIGATORIA:
            {perfil_activo}

            Idioma: {idioma} | Acento: {acento} | Tono: {tono}
            Tu salida debe ser ÚNICAMENTE un objeto JSON limpio.
        """
        
        # -- PROMPT PRINCIPAL: instrucciones completas --
        prompt = f"""
        --- OBJETIVO ---
        Diseña una estructura clara y jerárquica para enfocada en resonder la intencion de busqueda : '{title_base}'.
        La estructura debe estar optimizada y alineada con la intención de búsqueda.   


        --- INVESTIGACIÓN ---
        {contexto_completo_para_ia}

        --- REGLAS DE ESTRUCTURA ---

        1. H1 (H1 - 0)
            - Exactamente: '{title_base}'
            - No se repite ni se reformula.

        2. ALCANCE:

            -Se debe resolver la intención de búsqueda específica: '{title_base}'.
            -Prohibición de Alucinación: No inventes amplitud temática, eventos temporales o categorías que no resuelvan la intención de búsqueda.
            -Los encabezados deben ser relevantes y directamente vinculados a la intención de búsqueda.

        3. MÁXIMA CATEGORIZACIÓN Y EXCLUSIVIDAD (H2):
            
            - Cada H2 debe representar una categoría o subtema distinto solucionando la intención de búsqueda.
            - Prohibición de Redundancia: No repitas temas o subtemas en múltiples H2.
            - Secuencia Lógica: Los H2 deben seguir un orden lógico que guíe al lector a través del tema.   
            - REGLA ANTI-APÉNDICE: Prohibido crear secciones H2 o H3 para "Mapas", "Herramientas digitales", "Apps" o "Recursos". Estos elementos no son temas de lectura, son herramientas de apoyo.** 
        
        4. DESARROLLO SEMÁNTICO Y AGRUPACIÓN (H3):
            - Cada H3 debe expandir o detallar el H2 bajo el cual se encuentra.
            - Evita H3 que no aporten valor o que sean demasiado generales.
            - Si el titulo H2 contiene (LISTA) o (TABLA), no añadas H3 bajo ese H2.
             
        5. PROFUNDIDAD:
            - Prioriza la calidad y relevancia sobre la cantidad.
            - Si la intención de búsqueda se puede resolver en un número menor de H2/H3, hazlo.

        6. MULTIMEDIA (SUBORDINADA) (MAPAS, IMAGENES,  VIDEOS)
            - Tras cada H2 y H3:
            [MULTIMEDIA: TIPO | Descripción SEO]
            - Incluye multimedia SOLO cuando aporte al posicionamiento SEO o mejore la comprensión los titulos .
            - La multimedia refuerza la intención, nunca la define ni la amplía.
            - La descripción SEO debe ser breve y directa, enfocada en la función del multimedia en el contexto del encabezado.


        7. FORMATO DE SALIDA (OBLIGATORIO)
            - Formato: [H{{N}} - X.Y] Título del Encabezado
            - No uses Markdown.
            - No uses comillas, negritas, ni ningún otro formato a menos que sea en formato HTML.
            - No incluyas secciones de “Conclusión”, “FAQ” ni cierres editoriales.
            - Antes de entregar, revisa que la estructura cumple todas las reglas SEO para un posicionamiento óptimo, de lo contrario reestructura.

        --- FORMATO FINAL EXCLUSIVO ---
            Devuelve únicamente el objeto JSON:
            {{
                "structure_markdown": "Estructura optimizada siguiendo el formato [H{{N}} - X.Y] Título.\\n[MULTIMEDIA: TIPO | Descripción SEO]",
                "estimated_word_count": {estimated_word_count}
            }}            
        """
        
        
        
        # -- Llamada al modelo --
        response_json_str = self._llm_generate(
            prompt,
            system_message=system_message,
            temperature=0.4
        )

        print (response_json_str)

        # -- Limpieza y manejo de errores --
        try:
            return self.limpieza_extraccion_json(response_json_str)
        except Exception as e:
            return {
                "structure_markdown": f"[ERROR DE PARSEO CRÍTICO: {e} - Respuesta IA: {response_json_str[:200]}...]",
                "estimated_word_count": 0
            }
        

    # --- LOGICA PARA REGENERAR SOLAMENTE UNA UNICA PARTE DE LA ESTRUCTURA EN ESTE CASO TITULOS Y SUBTITULOS--- 
    def regenerar_titulos(self, 
        consolidated_text: str, 
        full_structure_markdown: str,
        section_to_regenerate: str,
        new_prompt: Optional[str] = None,
        idioma: str = "es", 
        acento: str = "neutral",
        tono: str = "profesional",
        **kwargs 
    ) -> List[str]:
        
        user_prompt_instruction = f"Instrucción de Edición/Regeneración Adicional: {new_prompt}\n" if new_prompt else ""
        
        prompt = f"""
        --- CONTEXTO COMPLETO DE REFERENCIA ---
        TEMA BASE: '{kwargs.get('main_title')}'
        CONTENIDO DE SCRAPING CONSOLIDADO: {consolidated_text[:2500]}... (Primeros 2500 caracteres como referencia de contexto)
        
        ESTRUCTURA ACTUAL COMPLETA DEL BLOG (Úsalo para CONTEXTO, pero NO LO REGENERES):
        ---
        {full_structure_markdown}
        ---

        --- TÍTULO/SUBTÍTULO A REGENERAR ---
        TÍTULO ACTUAL: '{section_to_regenerate}'
        
        --- INSTRUCCIONES CLAVE DE SALIDA ---
        1. **Formato OBLIGATORIO:** Debes devolver **SOLAMENTE** un array JSON de Python. Nada más.
        2. **Cantidad Estricta:** Genera **EXACTAMENTE 3 nuevas y mejoradas versiones** para el título/subtítulo: '{section_to_regenerate}'.
        3. **Contenido:** No incluyas el nivel jerárquico ([H2 - X.Y]) ni texto adicional. Solo las 3 frases.
        4. **PROHIBIDO:** No uses la clave "structure_markdown" o "estimated_word_count" en tu respuesta.
        
        EJEMPLO OBLIGATORIO DE SALIDA: ["Nueva Opción A", "Nueva Opción B", "Nueva Opción C"]
        
        {user_prompt_instruction}
        """
        
        system_msg = f'Eres un Estratega SEO y Redactor Creativo. Tu **ÚNICA Y EXCLUSIVA** tarea es generar **SOLO** un array JSON de 3 strings, siguiendo el formato estrictamente. No devuelvas ningún texto explicativo o formato JSON complejo, solo el array. Utiliza un **Tono {tono}** y **Acento {acento}** en idioma "{idioma}".'
        raw_response = self._llm_generate(prompt, system_msg, temperature=0.4)

        # 1. Limpieza de etiquetas Markdown (```json) y extracción del contenido crudo.
        clean_json_str = re.sub(r'```json\s*|```', '', raw_response, flags=re.IGNORECASE).strip()

        # 2. **Paso Clave:** Extraer SOLO el array JSON.
        array_match = re.search(r'(\[[\s\S]*?\])', clean_json_str, re.DOTALL)
        
        if array_match:
            json_to_parse = array_match.group(1).strip()
            
            # Proceso de parseo del array
            try:
                suggestions = json.loads(json_to_parse)
                if isinstance(suggestions, list) and all(isinstance(s, str) for s in suggestions):
                    return suggestions 
                else:
                    # La IA devolvió un array, pero no de strings o con formato incorrecto
                    print("ADVERTENCIA: Parseo exitoso, pero el resultado no es List[str] o es un array vacío.")
                    return [] 
                    
            except json.JSONDecodeError:
                # Lógica de corrección de comillas simples/dobles
                try:
                    corrected_json = json_to_parse.replace("'", '"')
                    suggestions = json.loads(corrected_json)
                    if isinstance(suggestions, list) and all(isinstance(s, str) for s in suggestions):
                        return suggestions
                    else:
                        # Parseo exitoso de JSON, pero no pasó la validación de tipos.
                        print("ADVERTENCIA: El JSON corregido se parseó, pero no es List[str].")
                        return []
                except (json.JSONDecodeError, ValueError) as e:
                    # Fallo irrecuperable de parseo JSON.
                    print(f"ERROR: Fallo de parseo JSON irrecuperable de array: {e}")
                    return [] 
            except Exception as e:
                # Cualquier otra excepción inesperada
                print(f"ERROR: Excepción inesperada durante el parseo de títulos: {e}")
                return [] 
        else:
            print("ADVERTENCIA: El LLM no devolvió un array JSON detectable ([...]).")
            return []

    #Regeneracion o generacion principal para los contenidos de los H 
    def generar_contenido_seccion (self, req: PeticionGeneracionContenido) -> Dict[str,Any]:
        """
        Logica para la regeneracion o generacion de contenido de una seccion especifica 
        (Maneja texto plano y el formato JSON para el modo de bloque libre).
        """

        # 1. Validacion de datos
        if not req.regenerate_data:
            raise HTTPException(status_code=400, detail="regenerate_data es requerido para la generacion de contenido.")
        
        # --- LÓGICA CLAVE: DETECCIÓN DEL MODO DE GENERACIÓN ---
        # Asumimos que 'section_type' se envía en la raíz del payload (req.section_type)
        is_block_generation = req.section_type == "content_generation_free"
        
        try: 
            # Campos base requeridos desde el frontes(Blog_generacion)
            section_title = req.regenerate_data['section_title']
            section_level = req.regenerate_data['section_level']
            full_structure_markdown = req.regenerate_data['full_structure_markdown']
            required_keywords: List[str] = req.regenerate_data.get('required_keywords', [])
            content_type: str = req.regenerate_data.get('content_type', 'parrafo_marrativo')
            context_data: Optional[str] = req.regenerate_data.get('context_data', None)
            estimated_word_count: Optional[int] = req.regenerate_data.get('estimated_word_count', 0)
            
            # Campo específico para el modo Bloque
            section_text: Optional[str] = req.regenerate_data.get('section_text', None) 

        except KeyError as e:
            raise HTTPException(status_code=400, detail=f"Falta el campo requerido en regenerate_data:{e}")
        
        # 2. Construccion de intrucciones dinamicas para el prompt
    
        # Lógica de Longitud (Texto Plano por defecto)
        length_instruction = (
            "El contenido debe ser **extremadamente conciso** y **solamente FUNDACIONAL** (a nivel de resumen o introducción). "
            "Desarrolla el tema de forma sintética, y **TERMINA INMEDIATAMENTE** el texto una vez que el concepto general de la sección haya sido cubierto. "
            "Bajo NINGUNA circunstancia introduzcas detalles que estén destinados a otras secciones de la estructura. Este es un punto CRÍTICO."
        )

        # Lógica de Keywords (Original)
        keyword_instruction = ""
        if required_keywords and isinstance(required_keywords, list):
            keywords_str = ', '.join(required_keywords)
            keyword_instruction = f"""
            INSTRUCCIÓN CLAVE DE SEO: Debes incluir las siguientes palabras clave en el texto: **{keywords_str}**.
            Es FUNDAMENTAL que te enfoques **únicamente** en el contexto de la sección '{section_title}' ({section_level}),
            evitando estrictamente temas y palabras clave que pertenezcan a otros H2/H3 de la estructura general para evitar la redundancia y el canibalismo semántico.
            """

        # Lógica de Contexto (Original)
        context_instruction = ""
        if context_data:
            context_instruction = (
                f"\n--- INSTRUCCIÓN CLAVE PARA PREVENCIÓN DE CANIBALISMO ---\n"
                f"La siguiente información describe temas que **YA ESTÁN CUBIERTOS** en otras partes del blog. "
                f"Asegúrate de que tu contenido **NO REPITA** o **REDUNDE** en las ideas listadas abajo, "
                f"sino que las complemente o las aborde desde un ángulo diferente.\n"
                f"{context_data}\n"
                f"--- FIN INSTRUCCIÓN CLAVE ---\n"
            )
        
        # --- LÓGICA DE FORMATO: CONDICIONAL (CORRECCIÓN CRÍTICA DE FORMATO) ---
        format_instruction = ""
        
        if is_block_generation:
            # Modo Bloque Libre: Instrucción CLAVE para forzar el JSON
            format_instruction = f"""
            INSTRUCCIÓN CRÍTICA DE FORMATO:
            DEBES DEVOLVER TU RESPUESTA **SOLAMENTE** COMO UN OBJETO JSON VÁLIDO.
            NO AÑADAS NINGÚN TEXTO EXPLICATIVO O INTRODUCCIÓN.
            El objeto JSON debe utilizar los títulos de las secciones listadas en el [BLOQUE DE ESTRUCTURA] como sus KEYS.
            """
            # Ajustamos la instrucción de longitud para el modo Bloque
            length_instruction = "Genera el contenido para todo el bloque, distribuyendo las palabras de manera lógica y coherente a través del H2 y sus H3s. Respeta el límite de palabras estimado del bloque."

        # Se mantiene el resto de la lógica de formatos si no es modo bloque:
        elif content_type == "lista_pasos":
            format_instruction = "El contenido debe ser una **lista numerada detallada** (1., 2., 3...) de pasos o instrucciones. Cada paso debe ser conciso, claro y estar en una línea separada."
        elif content_type == "lista_caracteristicas":
            format_instruction = "El contenido debe presentarse como una **lista con viñetas** (usando `*` o `-`) que enumere y describa brevemente ventajas, desventajas, características o elementos clave."
        elif content_type == "resumen_conciso":
            format_instruction = "El contenido debe ser un **párrafo único y conciso** (no más de 4-5 frases) que sirva como un resumen ejecutivo, una conclusión o un punto clave, con un lenguaje directo y persuasivo."
        elif content_type == "definicion_detallada":
            format_instruction = "El contenido debe iniciar con el término o frase, seguido de una definición clara y párrafos explicativos que profundicen en el concepto, su historia o su relevancia."
        elif content_type == "casos_texto":
            format_instruction = "El contenido debe enfocarse en proporcionar múltiples ejemplos o casos de uso prácticos que ilustren el tema. Cada ejemplo debe estar claramente separado, con su título en texto plano y su descripción en un párrafo."
        elif content_type == "comparacion_corta":
            format_instruction = "El contenido debe ser una comparación punto por punto entre 2 o 3 elementos clave (ej. Producto A vs. Producto B). Usa texto plano para los nombres de los elementos y viñetas para contrastar sus características de manera clara."
        elif content_type == "analisis_critico":
            format_instruction = "El contenido debe ser un **análisis estructurado en párrafos** con una introducción clara del problema o tema, un desarrollo del argumento central y una proyección o recomendación clara al final. Debe ser objetivo, sintético y basado en hechos."
        elif content_type == "pro_y_contra":
            format_instruction = "El contenido debe estar dividido en dos secciones claras: Pros (Ventajas) y Contras (Desventajas). Cada sección debe usar una lista con viñetas para enumerar y describir brevemente cada punto de manera equilibrada y separada. No uses negritas para los títulos de las secciones ni para los puntos."
        elif content_type == "datos_estadisticos":
            format_instruction = "El contenido debe enfocarse en presentar datos, cifras y estadísticas relevantes. Cada dato debe ser presentado en una línea separada, comenzando por el valor numérico, seguido de su explicación o contexto. No uses tablas, solo texto y listas."
        elif content_type == "mito_vs_realidad":
            format_instruction = "El contenido debe usar un formato de Mito vs. Realidad para desmentir conceptos erróneos. Cada punto debe tener una línea para el Mito y la siguiente línea para la Realidad (en formato de párrafo explicativo)."
        elif content_type == "linea_tiempo":
            format_instruction = "El contenido debe ser una línea de tiempo cronológica. Utiliza una lista numerada donde cada punto represente un hito o evento en la secuencia temporal, incluyendo el año o la fecha al inicio de cada punto."
        else: 
            format_instruction = (
                "Tu tarea es generar el contenido utilizando el formato que consideres más apropiado (párrafos y listas con viñetas) para el tema de la sección. "
                "PROHIBIDO: No utilices negritas, cursivas o subrayados. Solo texto plano."
            )

        # 3. Construcción del Prompt Principal (Ajuste para Modo Bloque)
        
        # Define el objetivo del prompt: o una sección simple, o el bloque completo.
        target_section_prompt = f"el contenido detallado para la sección con el título: '{section_title}', que pertenece al nivel de encabezado '{section_level}'."
        
        if is_block_generation:
            target_section_prompt = f"el contenido detallado para el bloque de estructura completa (H2 y sus H3s) definido en el [BLOQUE DE ESTRUCTURA]."

        # --- CORRECCIÓN DE SINTAXIS ---
        # Creamos la instrucción del bloque en una variable separada para evitar el SyntaxError
        block_instruction_str = ""
        if is_block_generation and section_text:
            # Usamos '\n' simple para el salto de línea, ya que estamos fuera de la f-string
            block_instruction_str = (
                f"BLOQUE DE ESTRUCTURA (Utiliza estos títulos como KEYS en el JSON):\n"
                f"{section_text}"
            )
        # --- FIN CORRECCIÓN DE SINTAXIS ---

        prompt = f"""
        Tu tarea **ÚNICA Y EXCLUSIVA** es generar {target_section_prompt}
        
        SOLO DEVUELVE EL TEXTO DEL CONTENIDO DE LA SECCIÓN SOLICITADA, SIN AÑADIR EL TÍTULO DE LA SECCIÓN NI NINGÚN OTRO ENCABEZADO (Aplica estrictamente para el modo Texto Plano).

        --- INSTRUCCIONES CRÍTICAS DE ALTA PRIORIDAD ---
        1. **ENFOQUE ESTRICTO Y ANTI-CANIBALISMO:** El contenido generado debe ser **DE NIVEL FUNDACIONAL Y SINTÉTICO** (como un resumen ejecutivo).
        Concéntrate **SOLAMENTE** en el tema específico de '{section_title}'.
        **ES OBLIGATORIO** que **NO** desarrolles ideas específicas (como ejemplos estatales o listas de pasos) que serán cubiertas en otros títulos de la estructura. 
        **UTILIZA** la información del bloque [INSTRUCCIÓN CLAVE PARA PREVENCIÓN DE CANIBALISMO] como tu regla principal de exclusión.
        
        2. **CONCISIÓN:** {length_instruction}
        
        3. **FORMATO:** {format_instruction}
        --- FIN INSTRUCCIONES CRÍTICAS ---
        
        El contenido debe ser en idioma '{req.idioma}' con acento '{req.acento}' y tono '{req.tono}'.

        {context_instruction} 
        
        {keyword_instruction}

        --- CONTEXTOS DE REFERENCIA (SOLO PARA INFORMACIÓN FACTUAL Y SEMÁNTICA) ---
        CONTEXTO DE LA ESTRUCTURA DEL BLOG (usa esto solo para tener el mapa completo):
        {full_structure_markdown}
        
        REFERENCIA DE CONTENIDO DEL SCRAPING (usa esto como fuente primaria de información y para asegurar la factualidad):
        {req.consolidated_content}

        {block_instruction_str}
        --- FIN CONTEXTOS DE REFERENCIA ---

        Asegúrate de que la salida respete estrictamente todas las INSTRUCCIONES CRÍTICAS provistas.
        """


        # 5. Llamada al LLM y Procesamiento
        generated_content = self._llm_generate(
            prompt=prompt,
            # Se refuerza el system_message
            system_message="Eres un escritor SEO profesional y un experto en condensación de contenido. Tu principal directriz es generar texto conciso, de alto valor narrativo y que respete estrictamente los límites temáticos para mantener el hilo conductor.",
            temperature=0.6
        )

        # 6. SIN VALIDACIÓN EN EL BACKEND - Únicamente limpieza mínima del texto.
        final_content_to_return = generated_content.strip()

        return {
            "generated_content": final_content_to_return,
            "success": str(True),
            "log": "Contenido generado exitosamente."
        }


    # -- GENERA EL CONTENIDO DE EL ESQUEMA DEL BLOG
    def generar_contenido_blog_libre(self, db: Session, req: models.AIAnalysisRequest) -> Dict[str, Any]:
        blog_id = req.blog_id

        # 1. Validación y Extracción de Datos
        if not req.regenerate_data:
            raise HTTPException(status_code=400, detail="El campo 'regenerate_data' es requerido para la generación completa.")

        try:
            data = req.regenerate_data
            section_title = data.get('section_title')
            section_level = data.get('section_level')
            
            full_structure_markdown = data.get('full_structure_markdown')
            section_to_generate_markdown = data.get('section_text') 
            estimated_word_count = data.get('estimated_word_count', 0)
            
            consolidated_content = getattr(req, "consolidated_content", None)
            
            # --- CARGA DESDE DB ---
            if blog_id and (not consolidated_content or not full_structure_markdown):
                try:
                    blog_uuid = UUID(str(blog_id))
                    scraping_data = db.query(Scraping).filter(Scraping.blog_id == blog_uuid).first()
                    
                    if scraping_data:
                        if not consolidated_content:
                            consolidated_content = getattr(scraping_data, 'consolidated_content', None)
                            if consolidated_content:
                                req.consolidated_content = consolidated_content
                        
                        if not full_structure_markdown:
                            full_structure_markdown = getattr(scraping_data, 'final_structure_markdown', None) 
                            if full_structure_markdown:
                                req.regenerate_data['full_structure_markdown'] = full_structure_markdown
                                data = req.regenerate_data
                except ValueError:
                    pass

            # RE-VALIDACIÓN
            if not req.consolidated_content: 
                raise HTTPException(status_code=400, detail="Error: 'consolidated_content' vacío.")
            if not data.get('full_structure_markdown'):
                raise HTTPException(status_code=400, detail="Error: 'full_structure_markdown' vacío.")
            if not all([section_title, data.get('full_structure_markdown'), section_to_generate_markdown]):
                raise ValueError("Faltan campos críticos en regenerate_data.")
            
            idioma = getattr(req, "idioma", None)
            acento = getattr(req, "acento", None)
            if not idioma or not acento:
                raise HTTPException(status_code=400, detail="Idioma/Acento requerido.")

        except (TypeError, ValueError) as e:
            raise HTTPException(status_code=400, detail=f"Error estructural: {e}")

        # 2. Preparación de Contexto (Historial) - OPTIMIZADO PARA EVITAR SATURACIÓN
        history_text = "No hay contenido previo."
        if isinstance(req.previous_content, list) and req.previous_content:
            # Enviamos solo la parte final del bloque anterior para dar continuidad sin duplicar tokens innecesarios
            ultimo_bloque = str(req.previous_content[-1])
            history_text = f"Conexión narrativa con el bloque anterior: ...{ultimo_bloque[-1000:]}"
        elif isinstance(req.previous_content, str) and req.previous_content:
            history_text = req.previous_content[-1000:]

        # 3. Lógica de Longitud
        instruccion_longitud = ""
        if estimated_word_count > 0:
            num_sections = len([
                line for line in full_structure_markdown.split('\n')
                if re.match(r'^\[H[1-9]\s*-\s*[0-9.]+\]', line.strip())
            ]) or 1
            target_avg_words = int(estimated_word_count / num_sections)
            instruccion_longitud = f"Referencia orientativa: **{target_avg_words} palabras**. Desarrolla con profundidad técnica/vivida."

       # 4. Extracción de Cabeceras para JSON
        headers_for_json = [section_title] 
        # Regex optimizada para capturar solo los títulos reales
        sub_headers_pattern = r'(\[H[0-9] - [0-9.]+\].*\s*(.+)$)|(^[#]{3,}\s*(.+)$)'
        sub_headers_raw = re.findall(sub_headers_pattern, section_to_generate_markdown, re.MULTILINE)
        
        sub_headers = []
        for match in sub_headers_raw:
            h_text = (match[1] or match[3]).strip()
            # Filtramos para que no entren etiquetas multimedia ni líneas vacías
            if h_text and "MULTIMEDIA" not in h_text.upper() and h_text != section_title:
                sub_headers.append(h_text)
        
        headers_for_json.extend(h for h in sub_headers if h not in headers_for_json)
        
        # Eliminamos posibles nulos o vacíos que rompen el prompt
        headers_for_json = [h for h in headers_for_json if h.strip()]
        
        json_schema_example = {header: f"[CONTENIDO PARA {header}]" for header in headers_for_json}
        json_schema_example_str = json.dumps(json_schema_example, indent=2, ensure_ascii=False)
        EDITORIAL_PROFILES = {
        "viajemos": """
            PROYECTO: VIAJEMOS (Voz de Experto que Acompaña)

            Voz Editorial:
            - No eres un folleto: Eres un viajero experimentado. Evita el tono institucional. Habla desde la recomendación real ("Te sugiero", "Lo mejor es").
            
            Enfoque de Datos SEO:
            - El valor está en el detalle: Si el Consolidated Content dice que algo cuesta 10 USD, no digas que es "barato", di que cuesta <strong>10 USD</strong>. La precisión es lo que genera confianza (E-E-A-T).
            
            Reglas de Oro del Tono:
            - Cero Relleno: Si vas a describir un lugar, usa detalles sensoriales (olores, colores, clima) mezclados con los datos técnicos.
            - Persuasión Útil: Cada párrafo debe ayudar al lector a tomar una decisión o imaginar el lugar.
            - Prohibido el lenguaje robótico: Sustituye "Miami ofrece diversas opciones" por "En Miami te vas a encontrar con...".
        """,

        "arriendo": """
            PROYECTO: ARRIENDO (Voz Educativa y Práctica)

            Voz Editorial:
            - Eres un asesor experto, neutral y didáctico. Tu objetivo es dar seguridad y claridad ante procesos legales o comerciales.

            Enfoque Narrativo:
            - Directo al grano: el lector busca soluciones, riesgos y consecuencias reales.
            - Define los escenarios de forma explicativa para que un ciudadano común los entienda sin esfuerzo.

            Reglas de Estilo:
            - Estilo limpio y profesional, sin storytelling emocional.
            - Prioriza la jerarquía de obligaciones, derechos y recomendaciones preventivas.
            - Transmite autoridad a través de la precisión y la sencillez.
            """,

        "guia_legal": """
            PROYECTO: GUÍA LEGAL (Voz de Autoridad y Confianza)

            Voz Editorial:
            - Eres un profesional del derecho que comunica con rigor pero con un lenguaje accesible para no-abogados. Transmites calma y respaldo experto.

            Enfoque Narrativo:
            - Informativo-legal puro. El texto debe responder con autoridad: qué es, cuándo aplica y qué acción debe tomar el lector.

            Reglas de Estilo:
            - Formal y profesional. Queda estrictamente prohibido el lenguaje promocional o emocional.
            - El valor reside en la exactitud de la norma y la claridad del procedimiento.
            - Cada párrafo debe reforzar la sensación de que el lector está en manos expertas.
        """
        }

        project_key = getattr(req, "project", "viajemos").lower()
        editorial_profile = EDITORIAL_PROFILES.get(project_key, EDITORIAL_PROFILES["viajemos"])



        # 5. Construcción del Prompt (El corazón de la simplificación) 
        system_message = f"""
            Eres un redactor experto en SEO, arquitectura semántica y edición editorial.
            Tu tarea es generar el contenido completo de la sección actual del blog
            respetando jerarquía temática, profundidad progresiva y coherencia global.

            ---

            CONTEXTO EDITORIAL
            - Título principal: {req.main_title}
            - Idioma: {idioma}
            - Acento: {acento}
            - Técnica: {req.tecnica}
            - Tono: {req.tono}
            - Autoridad SEO: Tu credibilidad se basa en la precisión. No ignores los datos técnicos del análisis; úsalos para dar peso al texto.

            Redacta con naturalidad humana, precisión informativa y control semántico.
            Evita texto genérico, redundante o reutilizable en otras secciones.

            ---

            ANÁLISIS INTERNO (NO MOSTRAR)
            1. Identifica el rol semántico de esta sección.
            2. Determina qué información pertenece EXCLUSIVAMENTE a este nivel.
            3. No anticipes, resumas ni dupliques contenido de secciones hermanas o hijas.

            ---

            REGLAS CRÍTICAS DE SALIDA
            - Devuelve únicamente un objeto JSON.
            - Las claves deben coincidir exactamente con los títulos/subtítulos.
            - Usa doble barra invertida (\\\\) para representar una barra literal (\).

            ---

            ### ALGORITMO DE ASIGNACIÓN Y CIERRE SEMÁNTICO (VERSIÓN SUPREMA)

            1. NIVEL H1 [Título 0] - ACTIVADOR DE INTENCIÓN:
            - Se debe incluir la intencion de busqueda principal del usuario.
            - Misión: Validar la intención de búsqueda (Search Intent).
            - Regla: PROHIBIDO usar datos técnicos, listas o precios del Consolidated Content. 
            - Esta parte es la introduccion de el articulo y debe ser narrativa y atractiva, preparando al lector para el desarrollo posterior.

            2. NIVEL H2 [Títulos 1, 2, 3...] - ORGANIZADOR TEMÁTICO:
            - Misión: Agrupar conceptos y establecer el contexto.
            - Regla: Detente antes de tocar el contenido detallado de los sub-bloques (H3).

            3. NIVEL H3 [Títulos 1.1, 2.1.1...] - UNIDAD DE VALOR Y EJECUCIÓN:
            - Misión: Entrega total de sustancia (E-E-A-T).
            - Obligación: Extraer CADA nombre propio, marca, precio o tecnicismo del 'Consolidated Content'. Si el dato existe, debe aparecer.

            ### PROTOCOLO DE CALIDAD Y RITMO EDITORIAL (ANTI-IA)

            - Variedad de Párrafos: Alterna párrafos de impacto con párrafos de desarrollo.
            - Escaneabilidad Semántica: Usa etiquetas HTML <strong> para resaltar datos duros (fechas, precios, lugares, cifras). No resaltes frases vacías; solo información útil para el SEO.
            - Naturalidad Narrativa: Evita muletillas de IA ("En este recorrido", "Prepárate para"). Usa afirmaciones directas de experto.
            - Cero Redundancia: Si un dato ya aparece en el 'Historial', queda estrictamente prohibido repetirlo. Pasa al siguiente detalle técnico.
        """
        

        # 6. Orquestación Inteligente por Lotes (Evita Saturación de Contexto)
        secciones_posteriores = full_structure_markdown.split(section_title)[-1][:500] if section_title in full_structure_markdown else ""
        
        # --- LÓGICA DE SEGMENTACIÓN ---
        MAX_HEADERS_PER_CALL = 2  # Bajamos a 2 para máxima estabilidad en servidor local
        header_chunks = [headers_for_json[i:i + MAX_HEADERS_PER_CALL] for i in range(0, len(headers_for_json), MAX_HEADERS_PER_CALL)]
        
        structured_content = {}
        contexto_acumulado = history_text 

        for index, chunk in enumerate(header_chunks):
            # FILTRADO DINÁMICO DEL CONSOLIDATED CONTENT
            # Solo pasamos las líneas del análisis que mencionen palabras clave de los títulos actuales
            lineas_analisis = req.consolidated_content.split('\n')
            analisis_relevante = []
            keywords = [word.lower() for h in chunk for word in h.split() if len(word) > 3]
            
            for linea in lineas_analisis:
                if any(key in linea.lower() for key in keywords):
                    analisis_relevante.append(linea)
            
            # Si el filtro es muy estricto, enviamos un fragmento base para no perder contexto
            fuente_datos_optimizada = "\n".join(analisis_relevante) if analisis_relevante else req.consolidated_content[:4000]

            chunk_schema = {header: f"[CONTENIDO PARA {header}]" for header in chunk}
            chunk_schema_str = json.dumps(chunk_schema, indent=2, ensure_ascii=False)

            prompt_lote = f"""
            ### PERFIL EDITORIAL: {project_key.upper()}
            {editorial_profile}

            ### CONTEXTO (Lote {index+1}/{len(header_chunks)})
            - **Sección:** {section_title}
            - **Títulos a redactar ahora:** {', '.join(chunk)}
            - **Historial:** {contexto_acumulado}

            ### DATOS ESPECÍFICOS DEL ANÁLISIS
            {fuente_datos_optimizada}

            ### TAREA
            Redacta el contenido para: {', '.join(chunk)}. 
            Usa los datos técnicos de la fuente de arriba. No inventes precios ni nombres.

            ### FORMATO JSON OBLIGATORIO
            ```json
            {chunk_schema_str}
            ```
            """

            try:
                # Bajamos max_tokens de salida para evitar que el servidor local se cuelgue procesando
                generated_response = self._llm_generate( 
                    prompt=prompt_lote,
                    system_message=system_message,
                    temperature=0.7, 
                    max_tokens=2500 
                )
                
                response_corrected = re.sub(r'(?<!\\)\\(?![ntrbvf/\\]|u[0-9a-fA-F]{4}|\"|\')', r'\\\\', generated_response)
                chunk_data = self.limpieza_extraccion_json(response_corrected)
                
                if isinstance(chunk_data, dict):
                    structured_content.update(chunk_data)
                    # Actualizamos historial para coherencia narrativa
                    ultimo_valor = list(chunk_data.values())[-1]
                    contexto_acumulado = f"Anteriormente escrito: {str(ultimo_valor)[-600:]}"
                
            except Exception as llm_e:
                print(f"Error lote {index+1}: {str(llm_e)}")
                # Si falla un lote, intentamos continuar con el siguiente en lugar de matar todo el proceso
                continue 

        # 6. Respuesta Final serializada
        return {
            "generated_content": json.dumps(structured_content), 
            "success": "True",
            "log": f"Generación finalizada. Se procesaron {len(structured_content)} de {len(headers_for_json)} secciones."
        }



    # --- AQUI ESTA LA LOGICA DE REGENERACION Y LIMPIEZA DEL JSON QUE DEVUELVE LA IA 
    def analisis_final_ia(
        self, 
        db: Session, 
        query: str, 
        title_base: str,
        categoria: str,
        idioma: str,
        tecnica: str,
        acento: str,
        tono: str, 
        blog_id: Optional[UUID] = None, 
        consolidated_text: Optional[str] = None, 
    ) -> Dict[str, Any]:
        
        # 1. OBTENER CONSOLIDATED CONTENT (Si aplica)
        if blog_id and not consolidated_text:
            # Se asume que 'Scraping' está importado
            scraping_result = db.query(Scraping).filter(Scraping.blog_id == blog_id).first()
            if scraping_result and scraping_result.consolidated_content:
                consolidated_text = scraping_result.consolidated_content
            
        if not consolidated_text:
            return {
                "structure_markdown": "[ERROR: No se pudo obtener el texto consolidado para la generación de la estructura.]",
                "estimated_word_count": 0
            }

        # ----------------------------------------------------------------------
        # 💥 CORRECCIÓN CRÍTICA: 2. OBTENER estimated_word_count DE LA TABLA BLOG
        # ----------------------------------------------------------------------
        if not blog_id:
            # Bloqueo si no hay ID para buscar la longitud
            raise HTTPException(status_code=400, detail="Blog ID es requerido para buscar 'estimated_word_count' en la DB.")

        # Se asume que 'Blog' está importado
        blog_entity = db.query(Blog).filter(Blog.id == blog_id).first()

        if not blog_entity:
            raise HTTPException(status_code=404, detail=f"No se encontró el Blog con ID: {blog_id}.")

        # Obtener el campo (MANDATO CUMPLIDO)
        estimated_word_count_from_db = getattr(blog_entity, 'estimated_word_count', None)
        
        # 3. VALIDACIÓN ESTRICTA (Según su mandato de que el dato siempre está lleno)
        if not estimated_word_count_from_db or not isinstance(estimated_word_count_from_db, int) or estimated_word_count_from_db <= 0:
            raise HTTPException(status_code=500, detail=f"El campo 'estimated_word_count' en el Blog ID {blog_id} es nulo o inválido.")
        
        # ----------------------------------------------------------------------
        # 4. LLAMADA A LA GENERACIÓN DE LA IA (Inyección del argumento)
        # ----------------------------------------------------------------------
        # NOTA: La función self.generar_estructura_seo_final DEBE haber sido actualizada
        # para aceptar 'estimated_word_count' como un argumento posicional/clave.
        analysis_result = self.generar_estructura_seo_final( 
            query=query,
            title_base=title_base,
            categoria=categoria,
            idioma=idioma,
            tecnica=tecnica,
            acento=acento,
            tono=tono,
            consolidated_text=consolidated_text,
            # ✅ ARGUMENTO REQUERIDO INYECTADO:
            estimated_word_count=estimated_word_count_from_db 
        )

        # ----------------------------------------------------------------------
        # 5. LIMPIEZA Y NORMALIZACIÓN DE LA ESTRUCTURA COMPLETA
        # ----------------------------------------------------------------------
        MARKDOWN_KEY = 'structure_markdown' 
        if MARKDOWN_KEY in analysis_result and isinstance(analysis_result[MARKDOWN_KEY], str):
            markdown_to_clean = analysis_result[MARKDOWN_KEY]
            markdown_to_clean = markdown_to_clean.replace('\r\n', '\n').replace('\r', '\n')
            markdown_to_clean = re.sub(r'[^\S\n]+', ' ', markdown_to_clean) 
            markdown_to_clean = '\n'.join([line.strip() for line in markdown_to_clean.split('\n') if line.strip()])
            markdown_to_clean = markdown_to_clean.strip()
            analysis_result[MARKDOWN_KEY] = markdown_to_clean
                
        # ----------------------------------------------------------------------
        # 6. PERSISTENCIA DE DATOS
        # ----------------------------------------------------------------------
        if blog_id:
            # Persistencia de la Estructura en la tabla BLOGS
            try:
                actualizar_estructura_blog(db, blog_id, analysis_result) 
            except Exception as e:
                print(f"ERROR DE PERSISTENCIA: Fallo al guardar la estructura en la tabla 'blogs': {e}")
                
            # Persistencia del Conteo de Palabras en la tabla SCRAPING
            if 'estimated_word_count' in analysis_result:
                try:
                    word_count_str = analysis_result['estimated_word_count']
                    word_count_int = int(word_count_str) if str(word_count_str).isdigit() else None
                    
                    if word_count_int is not None:
                        datos_para_scraping = {
                            "estimated_word_count_ai": word_count_int 
                        }
                        actualizar_o_crear_resultado_scraping(db, blog_id, datos_para_scraping)
                        print(f"estimated_word_count_ai persistido con éxito en Scraping para Blog ID: {blog_id}")
                except Exception as e:
                    print(f"ERROR DE PERSISTENCIA: Fallo al guardar estimated_word_count en la tabla 'scraping': {e}")
                        
        # ----------------------------------------------------------------------
        # 7. DEVOLVER LA RESPUESTA
        # ----------------------------------------------------------------------
        return analysis_result



# --- 2. CLASE ContentExtractor: Lógica de Scraping y Fallback ---
class ContentExtractor:
    """
    Clase encargada de la descarga, limpieza, extracción estructurada (Plan A/B)
    y la gestión de la lógica de fallback de scraping.
    """
    # --- CONSTANTES Y PATRONES DE RUIDO ---
    EXCLUDED_HEADINGS = [
        'contenido relacionado', 'principales noticias', 'no te lo pierdas', 
        'lecturas más populares', 'más información', 'comentarios', 
        'temas relacionados', 'navegación', 'créditos', 'síguenos', 
        'te puede interesar', 'leer más', 'share', 'siguiente', 'anterior', 
        'ver también', 'suscríbete', 'regístrate', 'cierra sesión', 'publicidad', 
        'productos recomendados', 'Contenido relacionado', 'Temas relacionados',
        'Lecturas más populares', 'patrocinado','opinión', 'redes sociales', 'directorio', 'aviso de privacidad', 
        'términos y condiciones', 'nuestras redes', 'miembro del grupo de diarios de américa',
        'código de ética', 'consultas', 'newsletters', 'juegos', 'podcast', 
        'videos', 'publicidad', 'newsletter', 'lo más visto', 'lo más leído',
        'contacto', 'acerca de', 'autor','te recomendamos', 'también te puede interesar', 'otras noticias', 
        'más sobre este tema', 'otras historias', 'otras noticias de', 
        'historias relacionadas', 'otras historias de', 'lo que te podría interesar',
        'noticias de america latina', 'noticias internacional', 'lo más visto',
    ]

    # ---PATRONES A EXCLUIR ---
    NOISE_PATTERNS = re.compile(
        r'Fuente\s+de\s+la\s+imagen[,\s]+(Getty\s+Images|Alamy)\s*|'
        r'Información\s+del\s+artículo\s+Autor[,\s]+Redacción\s+Título\s+del\s+autor[,\s]+BBC\s+Travel\s+[\d\s\w]+\s*|'
        r'Getty\s+Images\.?\s*|'
        r'Agencia\s+EFE\s*|'
        r'\s*Lo\s+mejor\s*:\s*'
        r'\s{2,}|\n|\t|[\xa0\ufeff]+', 
        re.UNICODE | re.I | re.M 
    )

    def __init__(self):
        self.ua = UserAgent()
        self.scraper = cloudscraper.create_scraper(
            browser={
                'browser': 'chrome',
                'platform': 'windows',
                'desktop': True
            }
        )

    @staticmethod
    def is_relevant_src(src: str | None) -> bool:
        if not src or len(src) < 10 or src.startswith('data:image') or src.startswith('#'):
            return False
        if any(s in src.lower() for s in ['adserve', 'sponsors', 'tracker', 'g-recaptcha', 'pixel', 'comments-area']):
            return False
        return True


    @staticmethod
    def clean_text(text: str) -> str:
        """Elimina el ruido patronizado (ej. créditos de imagen) y recorta los espacios extra del texto extraído."""
        if not text: return ""
        return ContentExtractor.NOISE_PATTERNS.sub(' ', text).strip()

    @staticmethod
    def _get_media_info(tag: Tag) -> Dict[str, str] | None: 
        """
        Extrae la URL de origen y el texto descriptivo (alt/caption) de un elemento multimedia, 
        manejando atributos de carga perezosa, miniaturas de YouTube y contenedores genéricos.
        """

        PRIMARY_MEDIA_TAGS = ['img', 'picture', 'iframe', 'video', 'figure', 'source']
        
        # Lista ampliada de atributos de Lazy-Loading y URLs de fondo
        LAZY_ATTRIBUTES = [
            'data-src', 'data-original', 'data-url', 'data-lazy-src', 'data-image-src', 
            'srcset', 'data-srcset', 'data-iframe-src', 'data-lazyload', 'data-large-file',
            'data-bg-src', 'data-bg', 'data-srcset-mobile' 
        ]
        
        # --- FILTRO DE ENTRADA PERMISIVO ---
        is_media_tag = tag.name in PRIMARY_MEDIA_TAGS
        # Detección de la miniatura precargada de YouTube (clase específica)
        is_youtube_div = tag.name == 'div' and 'ytp-cued-thumbnail-overlay' in tag.get('class', [])
        # Detección de contenedores genéricos con atributos de carga perezosa
        is_lazy_container = tag.name in ['div', 'span', 'section'] and any(tag.get(attr) for attr in LAZY_ATTRIBUTES)

        if not is_media_tag and not is_youtube_div and not is_lazy_container:
            return None
        
        src = tag.get('src')
        
        # --------------------------------------------------------------------------
        # 1. LÓGICA DE CARGA PEREZOSA (data-*)
        # --------------------------------------------------------------------------
        if not src or (is_lazy_container and not src): 
            for attr in LAZY_ATTRIBUTES:
                if tag.get(attr):
                    src_value = tag.get(attr)
                    
                    # Manejo de srcset/data-srcset
                    if attr in ['srcset', 'data-srcset', 'data-srcset-mobile'] and src_value:
                        try:
                            last_pair = src_value.split(',')[-1].strip()
                            src = last_pair.split(' ')[0]
                        except:
                            src = None
                    else: 
                        src = src_value
                    if src: break
            
        # --------------------------------------------------------------------------
        # 2. DETECCIÓN DE IMÁGENES/VIDEOS EN ATRIBUTO 'STYLE' (Fallo común en blogs)
        # --------------------------------------------------------------------------
        if not src and tag.get('style'):
            style_attr = tag.get('style')
            
            # Búsqueda de background-image: url(...)
            style_match_img = re.search(r'background-image:\s*url\s*\(["\']?(.+?)["\']?\)', style_attr, re.I)
            if style_match_img:
                src = style_match_img.group(1)
            
            # Detección de miniatura de YouTube por estilo si falló la clase
            elif 'ytp-cued-thumbnail-overlay-image' in tag.get('class', []):
                match_url = re.search(r'url\("?(.+?)"?\)', style_attr)
                if match_url:
                    thumb_url = match_url.group(1).replace('"', '')
                    id_match = re.search(r'/vi/([a-zA-Z0-9_-]+)/', thumb_url)
                    if id_match:
                        src = f"https://www.youtube.com/embed/{id_match.group(1)}"


        # --------------------------------------------------------------------------
        # 3. DETECCIÓN POR CLASE YOUTUBE (Si el src aún está vacío y es el div)
        # --------------------------------------------------------------------------
        if is_youtube_div and not src:
            try:
                image_div = tag.find('div', class_='ytp-cued-thumbnail-overlay-image')
                if image_div:
                    style_attr = image_div.get('style', '')
                    match_url = re.search(r'url\("?(.+?)"?\)', style_attr)
                    
                    if match_url:
                        thumb_url = match_url.group(1).replace('"', '')
                        id_match = re.search(r'/vi/([a-zA-Z0-9_-]+)/', thumb_url)
                        if id_match:
                            # Creamos la URL de embed del video
                            src = f"https://www.youtube.com/embed/{id_match.group(1)}"
            except:
                pass 


        # --------------------------------------------------------------------------
        # 4. LÓGICA DE DETECCIÓN FINAL Y CLASIFICACIÓN
        # --------------------------------------------------------------------------
        if src and ContentExtractor.is_relevant_src(src):
            
            media_type = 'Imagen'; alt_text = tag.get('alt', '')
            
            # Clasificación para tags multimedia primarios
            if tag.name in ['iframe'] and any(keyword in src for keyword in ['youtube.com', 'youtu.be', 'vimeo.com', 'maps.google.com']): 
                media_type = 'Video/Mapa' 
                alt_text = tag.get('title') or alt_text
            elif tag.name == 'video' or any(ext in src.lower() for ext in ['.mp4', '.webm', '.ogg']): 
                media_type = 'Video'
            
            # Clasificación para tags contenedores genéricos
            elif tag.name in ['div', 'span', 'section'] or is_lazy_container:
                if any(ext in src.lower() for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg']):
                    media_type = 'Imagen (Lazy/CSS)'
                elif any(ext in src.lower() for ext in ['youtube.com', 'youtu.be', 'vimeo.com']):
                    media_type = 'Video (Lazy/CSS)'
                else:
                    media_type = 'Otro Multimedia'


            # Intenta obtener el pie de foto (caption)
            caption_text = ''
            search_scope = tag if tag.name == 'figure' else tag.find_parent('figure')
            
            if search_scope:
                    caption_tag = search_scope.find(
                        ['figcaption', 'span', 'p'], 
                        class_=re.compile(r'caption|pie-de-foto|fig-caption|description|credit', re.I)
                    ) 
                    caption_text = caption_tag.get_text(strip=True) if caption_tag else ''

            source_text = caption_text if caption_text else alt_text if alt_text else f"Multimedia de tipo {media_type}"
            
            return {'type': media_type, 'url': src, 'description': source_text, 'alt': alt_text, 'caption': caption_text} 
        
        return None

    @staticmethod
    def get_article_main_heading(soup: BeautifulSoup) -> str:
        """Busca y retorna el título principal del artículo, priorizando la etiqueta H1 y selectores comunes de titular."""
        
        h1_tag = soup.find('h1', class_=lambda c: c is None or 'logo' not in c.lower() and 'brand' not in c.lower() and 'site-title' not in c.lower())
        
        if h1_tag and h1_tag.get_text(strip=True) and len(h1_tag.get_text(strip=True)) > 10:
            return h1_tag.get_text(strip=True)
        
        # Búsqueda basada en clases comunes de título
        title_classes = re.compile(r'main-title|article-title|entry-title|post-title|headline-text', re.I)
        main_div = soup.find(['div', 'span'], class_=title_classes)
        
        if main_div and main_div.get_text(strip=True) and len(main_div.get_text(strip=True)) > 10:
            return main_div.get_text(strip=True)
            
        title_tag = soup.find("title")
        
        return title_tag.get_text(strip=True) if title_tag else "Contenido Principal del Artículo"


    @staticmethod
    def _get_content_area(soup: BeautifulSoup, mode: str) -> Union[Tag, BeautifulSoup]:
        """
        Identifica el área principal del contenido usando selectores simples (Plan A) 
        y luego heurística de densidad (Plan B/Modo Robusto). 
        Aplica limpieza fina interna para eliminar ruido anidado.
        """
        temp_content_area = None
        
        # ----------------------------------------------------------------------
        # A. DETECCIÓN SIMPLE (Plan A: Primer intento rápido y eficiente)
        # ----------------------------------------------------------------------
        if mode == 'simple':
            # Selectores de muy alta confianza
            simple_selectors = ['div[itemprop*="articleBody"]', '.entry-content', 'article.post-content', 'article', 'main']
            
            for selector in simple_selectors:
                area = soup.select_one(selector)
                # Debe tener suficiente texto para ser un artículo real
                if area and len(area.get_text(strip=True)) > 500: #<-- En caso de necesitar menos palabras para el scrapping reducirlo
                    temp_content_area = area
                    break

        # ----------------------------------------------------------------------
        # B. DETECCIÓN ROBUSTA (Plan B: Heurística de Densidad y Selectores Agresivos)
        # Se ejecuta si el modo es 'robust' O si el Plan A falló.
        # ----------------------------------------------------------------------
        
        if mode == 'robust' or (mode == 'simple' and not temp_content_area):
            
            # Selectores de contenedores de artículo 
            article_container_selectors = [
                'div[itemprop*="articleBody"]', '.entry-content', '.post-content', 
                '.article-body', '.post-body', '.article-main-content', '.td-post-content', 
                '.post-inner', '.content-wrap', '.single-post-content', 
                
                # Selectores Agresivos
                '[class*="content"]', '[class*="single"]', '.post', '.article',
                '.main-content', '#main-content-area', '.main-area',
                
                'article', 'main', '#content', '#primary', '#main-content', 
            ]
            
            best_area = temp_content_area
            max_p_count = 0
            
            # Si venimos de un fallo en el modo simple, reiniciamos max_p_count
            if not best_area:
                 max_p_count = 0

            # Iteramos sobre todos los candidatos para encontrar el mejor por DENSIDAD
            for selector in article_container_selectors:
                for area in soup.select(selector): 
                    # Contamos párrafos (el marcador clave de contenido principal)
                    p_count = len(area.find_all('p', limit=10)) 
                    text_len = len(area.get_text(strip=True))
                    
                    # Criterio de Selección: Debe ser grande, tener al menos 3 párrafos y más que el candidato actual.
                    if text_len > 300 and p_count >= 3:
                         if p_count > max_p_count:
                            max_p_count = p_count
                            best_area = area
            
            temp_content_area = best_area
        
        # ----------------------------------------------------------------------
        # C. LÓGICA DE LIMPIEZA INTERNA (LISTA BLANCA + EXCLUSIÓN DE RUIDO)
        # ----------------------------------------------------------------------

        if temp_content_area:
            # Lista de etiquetas esenciales que SI deben sobrevivir 
            essential_tags = ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'li', 'a', 'strong', 'em', 'blockquote', 'img', 'figure', 'ul', 'ol', 'video', 'table', 'span', 'br', 'iframe']
            
            for element in temp_content_area.find_all(True):
                is_noise = False
                
                # Criterio C.1: Eliminación por Lista Blanca y Vacío (Protege el contenido esencial)
                if element.name not in essential_tags:
                    # Busca si el elemento tiene un hijo multimedia 
                    has_media_child = element.find(['img', 'figure', 'iframe', 'video', 'picture'], recursive=False)
                    # Si no es un tag esencial y NO tiene texto significativo
                    if len(element.get_text(strip=True)) < 50 and not has_media_child: 
                        is_noise = True

                # Criterio C.2: Heurística de Contenedor de Ruido (Clases y Densidad de Enlaces)
                if element.name in ['div', 'section', 'aside']:
                    # 1. Clases de ruido conocidas (Filtro AMPLIO)
                    if any(c in element.get('class', []) for c in [
                        'widget', 'promo-box', 'related-posts', 'guide-links', 
                        'reviews-section', 'author-box', 'social-media', 'share-bar', 
                        'elementor-widget', 'ad-container', 'post-nav', 'links-list', 
                        'more-stories', 'paywall-block', 'sub-header', 'footer-content', 'byline-item'
                    ]):
                         is_noise = True
                         
                    # 2. Heurísticas de densidad de enlaces (Detecta listados de links/publicidad)
                    num_p = len(element.find_all('p', recursive=False))
                    num_a = len(element.find_all('a'))
                    text_len = len(element.get_text(strip=True))

                    # Regla: Si tiene muy pocos párrafos, muchos enlaces y poco texto total.
                    if num_p < 3 and num_a > 3 and num_a / (len(element.find_all(True)) or 1) > 0.3 and text_len < 500:
                        is_noise = True
                        
                # Criterio C.3: HEURÍSTICA DE CÓDIGO JS Y BOTONES DE ACCIÓN
                element_text = element.get_text(strip=True)
                
                # Detecta tags de form, input o placeholders de JS/React
                if element.name in ['input', 'textarea', 'select', 'form'] or '{{' in element_text or '}}' in element_text: 
                     is_noise = True
                
                # Detecta botones o links de acción
                elif element.name in ['button', 'a'] and any(keyword in element_text for keyword in ['Borrar', 'Selecciona', 'Reservar', 'Comprar', 'Agregar', 'Idioma', 'Moneda', 'Opciones', 'Destino', 'Ver más', 'Buscar', 'Suscribir', 'Newsletter', 'Publicidad']):
                    is_noise = True
                
                if is_noise:
                    element.decompose() # Elimina el elemento
            
            # 4. Heurística Final para eliminar bloques de Related Posts/Links al final
            for last_tag in temp_content_area.find_all(recursive=False)[-3:]:
                if last_tag.name in ['div', 'section'] and (
                    len(last_tag.get_text(strip=True)) < 500 and len(last_tag.find_all('a')) > 5
                ):
                    last_tag.decompose()

        # ----------------------------------------------------------------------
        # D. FALLBACK FINAL
        # ----------------------------------------------------------------------
        return temp_content_area or soup.find('body') or soup

    @staticmethod
    def group_content_by_headings(soup: BeautifulSoup, mode: str) -> List[Dict[str, Any]]: 
        """Procesa el HTML de la página, agrupando texto y elementos multimedia bajo los encabezados (H2/H3) detectados para crear bloques estructurados de contenido."""
        blocks = []
        current_heading: Optional[str] = None 
        current_content: List[str] = []
        current_media: List[Dict[str, str]] = [] 
        
        content_area = ContentExtractor._get_content_area(soup, mode)
        article_title = ContentExtractor.get_article_main_heading(soup)

        # Función auxiliar interna para guardar el bloque actual
        def save_current_block():
            heading_to_save = current_heading if current_heading is not None else "Introducción/Pre-H2"
            if current_content or current_media:
                safe_content = [str(c) for c in current_content if c is not None]
                cleaned_content = ContentExtractor.clean_text(" ".join(safe_content))
                deduplicated_media = {item.get('url'): item for item in current_media if item.get('url')} 
                if cleaned_content or deduplicated_media:
                    blocks.append({
                        "heading": heading_to_save,
                        "content": cleaned_content,
                        "media_elements": list(deduplicated_media.values())
                    })
        
        # Tags a buscar para la división y el contenido
        elements_to_find = ['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'blockquote', 'div', 'section', 'img', 'figure', 'iframe', 'video', 'picture', 'span']
        elements = content_area.find_all(elements_to_find, recursive=True)
        
        HEADING_CLASSES = re.compile(r'title|subtitle|headline|subhead|h-?[234]|h[234]-?style', re.I)
        LOW_CONFIDENCE_CLASSES = re.compile(r'caption|credit|footer|ad|widget|photo|image|media-title|social-share|figure-title|byline|author|metadata', re.I)
        
        for tag in elements:
            # Lógica para identificar si el elemento es un divisor de encabezado (H2/H3/Div con clase de título)
            tag_name = tag.name.lower()
            text_content = tag.get_text(strip=True, separator=' ')
            
            if tag.get('aria-hidden') == 'true' or tag.get('role') == 'presentation' or (not text_content and tag_name not in ['img', 'figure', 'picture', 'iframe']):
                continue

            is_heading_divisor = False
            
            # 1. Identificación del Divisor (H1, H2, H3/H4 bajo ciertas condiciones, y divs/spans con clases de encabezado)
            if tag_name in ['img', 'figure', 'picture', 'iframe', 'video']: is_heading_divisor = False
            elif tag_name in ['h1', 'h2']: is_heading_divisor = True
            elif tag_name in ['h3', 'h4']:
                is_heading_divisor = True
            elif tag_name in ['div', 'span', 'section']:
                if tag.get('role') == 'heading' and tag.get('aria-level') in ['1', '2', '3', '4']: is_heading_divisor = True
                elif tag.has_attr('class'):
                    class_str = " ".join(tag.get('class', []))
                    class_match = HEADING_CLASSES.search(class_str)
                    if class_match and len(text_content) > 10:
                        low_conf_match = LOW_CONFIDENCE_CLASSES.search(class_str)
                        if not low_conf_match:
                            has_media_child = tag.find(['img', 'figure', 'video', 'iframe'])
                            if not has_media_child or len(tag.contents) > 2: is_heading_divisor = True
            
            # 2. Manejo de Encabezado: Guarda el bloque anterior e inicia uno nuevo
            if is_heading_divisor:
                
                # FILTRO : Títulos muy cortos o conocidos como ruido visual
                if len(text_content) < 5 or any(exc in text_content.lower() for exc in ['pie de foto', 'foto:', 'imagen de', 'ver galeria', 'crédito']): continue
                

                # INICIO DE FILTRO UNIVERSAL CONTRA EL RUIDO DE LAS PAGINAS
                # 1. Chequeo de calidad del bloque ANTERIOR antes de guardarlo.
                if current_heading is not None and current_content:
                    
                    texto_consolidado = " ".join(current_content)
                    len_texto = len(texto_consolidado)
                    heading_strip = current_heading.strip()
                    
                    # FILTRO 1: DUPLICACIÓN HEADER-CONTENIDO EXTENDIDA 
                    if len_texto < 1500 and texto_consolidado.startswith(heading_strip): 
                        current_heading = text_content
                        current_content = []
                        current_media = []
                        continue 

                    # FILTRO 2: DENSIDAD DE ENLACES ADAPTATIVA MÁS AGRESIVA 
                    link_count = texto_consolidado.lower().count('http') + texto_consolidado.lower().count('www.')
                    
                    # Si el bloque es PEQUEÑO (< 500 chars) Y DENSO EN ENLACES (>= 3), es ruido.
                    if len_texto < 500 and link_count >= 3: 
                        current_heading = text_content
                        current_content = []
                        current_media = []
                        continue 
                        
                # FIN DE FILTROS UNIVERSALES CONTRA EL RUIDO EN LAS PAGINAS

                if current_heading is None or text_content.strip() != current_heading.strip():
                    save_current_block()
                    current_heading = text_content; current_content = []; current_media = []
                
            # 3. Manejo de Contenido Multimedia
            media_info = ContentExtractor._get_media_info(tag) 
            if media_info: current_media.append(media_info) 
                
            # 4. Manejo de Contenido Textual
            is_content_tag = tag_name in ['p', 'ul', 'ol', 'blockquote']
            
            # INCULYE LOS DIVS Y SECTION EN EL CONTENIDO DE LAS PAGINAS 
            # Si no es un encabezado y tiene texto sustancial, es contenido.
            if tag_name in ['div', 'section'] and not is_heading_divisor and len(text_content) > 50:
                is_content_tag = True
                
            if tag_name == 'span' and not is_heading_divisor and len(text_content) > 100: is_content_tag = True
            
            if is_content_tag and text_content and len(text_content) > 10: 
                if not is_heading_divisor: current_content.append(text_content)
                    
        save_current_block()
        
        if blocks and blocks[0]['heading'] == "Introducción/Pre-H2":
            blocks[0]['heading'] = article_title 
        
        final_blocks = []
        
        MIN_TEXT_LENGTH_CHARS = 50 
        
        for block in blocks:
            heading_lower = block['heading'].lower().strip()
            is_irrelevant_heading = any(exc in heading_lower for exc in ContentExtractor.EXCLUDED_HEADINGS)
            if is_irrelevant_heading: continue
            content_len = len(block['content']); has_media = bool(block['media_elements'])
            is_substantial = (content_len >= MIN_TEXT_LENGTH_CHARS) or (content_len > 0 and has_media) 
            if is_substantial: final_blocks.append(block)
                
        return final_blocks


    def fetch_webpage(self, url: str) -> tuple[str, str, BeautifulSoup]:
        """
        Descarga la página web usando cloudscraper.
        """
        try:
            # Desactivamos los warnings molestos de SSL en la consola
            urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
            
            headers = {'User-Agent': self.ua.random}
            
            # EL CAMBIO CLAVE:
            # Usamos el scraper pero SIN el parámetro verify=False dentro del .get()
            # Cloudscraper ya se encarga de negociar el SSL correctamente.
            response = self.scraper.get(url, headers=headers, timeout=15)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, "html.parser")
            
            # --- Tu lógica de limpieza original (Mantenida intacta) ---
            for tag in soup(["script", "style", "form"]): 
                tag.decompose()

            irrelevant_selectors = [ 
                'header', 'footer', 'nav', 'aside', '[role*="complementary"]', 
                '#sidebar', '#footer', '#header', '#top-menu', '[data-testid*="footer"]', 
                '.hero-section', '#skip-link', 'a[href="#main-content"]', 
                '[class*="ad-"]', '[id*="ad-"]', '[class*="advert"]', '[class*="sponsor"]', 
                '[id^="ezoic-"]', '#newsletter-form', '[class*="paywall"]', 
                '[class*="cookie"]', '[id*="consent"]', '[id*="onetrust"]', 
                '.c-cookie-banner', '.paywall-meter', '.g-ui-layer', 
                '[class*="related"]', '[class*="suggested"]', '[class*="recommend"]', 
                '[class*="next-story"]', '.article-meta', '.byline', '.metadata', 
                '.author-info', '.date-info', '[class*="author-box"]', 
                '[class*="share-bar"]', '[class*="social-media"]', '#comments', 
                '.article-comments', '.SocialShare', '.reviews-section', 
                '.guide-links', '.related-cities', '.language-selector', 
                '.currency-selector', '[class*="selector"]', '[class*="options"]',
                'input', 'form', 'button'
            ]
            
            for selector in irrelevant_selectors:
                for element in soup.select(selector):
                    try: element.decompose()
                    except: pass 
            
            # Extraemos texto limpio
            text = trafilatura.extract(response.text)
            if not text:
                text = re.sub(r"\s+", " ", soup.get_text()).strip()

            title = soup.find("title")
            title_text = title.get_text().strip() if title else "Sin título"

            return title_text, text, soup

        except Exception as e:
            logging.error(f"Error en fetch_webpage para {url}: {e}")
            return "Error de Scrapeo", "", BeautifulSoup("", "html.parser")
        

    def _scrape_with_fallback(self, soup: BeautifulSoup) -> List[Dict[str, Any]]:
        """
        Ejecuta la extracción de contenido usando el modo 'simple' (Plan A). Si el resultado 
        es insuficiente (por debajo de umbrales mínimos de contenido o bloques), utiliza el modo 'robust' (Plan B) como fallback.
        """
        MIN_CONTENT_CHARS = 300  
        MIN_BLOCKS = 5           
        
        # 1. Intento Sencillo (Plan A)
        blocks_simple = ContentExtractor.group_content_by_headings(soup, mode='simple') 
        
        # Evaluación del Plan A
        total_content_len_simple = sum(len(b.get('content', '')) for b in blocks_simple)
        is_simple_successful = total_content_len_simple >= MIN_CONTENT_CHARS and len(blocks_simple) >= MIN_BLOCKS
        
        if is_simple_successful:
            return blocks_simple
            
        # 2. Fallback Robusto (Plan B)
        blocks_robust = ContentExtractor.group_content_by_headings(soup, mode='robust')
        
        return blocks_robust
    
    


# --- 3. CLASE AnalysisOrchestrator: Coordinación del Flujo ---

class AnalysisOrchestrator:
    """
    Clase que coordina la extracción, el análisis de keywords y las llamadas a la IA.
    Contiene la lógica del flujo de 'execute_scraping'.
    """

    def __init__(self, ai_service: AIService, extractor: ContentExtractor, db: Session, blog_id: UUID):
        """
        Inicializa el orquestador con servicios de IA y Extracción, y las dependencias de DB.
        """
        self.ai_service = ai_service
        self.extractor = extractor
        self.db = db # <-- GUARDAR LA SESIÓN DB
        self.blog_id = blog_id # <-- GUARDAR EL ID DEL BLOG

    
    def _aggressive_text_fallback(self, soup: BeautifulSoup, article_heading: str) -> List[Dict[str, Any]]:
        """
        Plan C: Intento de extracción de texto agresiva cuando la división estructural falla.
        Devuelve el contenido en un solo bloque.
        """
        
        content_area: Tag | BeautifulSoup = ContentExtractor._get_content_area(soup, mode='robust') 
        
        text_elements = content_area.find_all(['p', 'h2', 'h3', 'li', 'p','h4', 'h5', 'h6','strong'])
        
        aggressive_text = []
        for elem in text_elements:
            text = ContentExtractor.clean_text(elem.get_text(strip=True))
            if len(text) >= 3: 
                text_lower = text.lower()
                if 'skip to' not in text_lower and 'search' not in text_lower:
                    aggressive_text.append(text)

        fallback_text = "\n\n".join(aggressive_text)

        if len(fallback_text) < 5:
            try:
                # Fallback ultra-agresivo (todo el texto del cuerpo)
                content_to_extract = soup.body if soup.body else soup
                raw_text_ultra = content_to_extract.get_text(separator=' ', strip=True) 
                fallback_text = ContentExtractor.clean_text(raw_text_ultra)
                if len(fallback_text) == 0: raise ValueError("Contenido del documento es CERO.")
            except Exception:
                return [] 

        return [{"heading": article_heading, "content": fallback_text, "media_elements": []}]

    def execute_scraping(self, req: models.ScrapeRequest) -> Generator[str, None, None]:
        """
        Ejecuta el flujo completo de scraping y envía eventos de progreso al frontend.
        """
        # 1. Extraer URLs
        urls = [u.url.strip() for u in req.urls if u.url.strip()][:req.num_results]
        analisis_contexto = req.title_base if req.title_base else "Contenido relevante"

        all_results = []
        consolidated_text = "" 
        log = [f"Iniciando scraping de URLs a las {datetime.now().strftime('%H:%M:%S')}"]
        
        print(f"Iniciando scraping de {len(urls)} URLs")

        for i, url in enumerate(urls, 1):
            # 📢 ENVIAR AL FRONTEND: "Empezando URL X"
            msg_inicio = f"Procesando URL {i} de {len(urls)}: {url}"
            print(msg_inicio)
            yield f"data: {msg_inicio}\n\n"

            # 1. Scraping y Limpieza
            title, full_text, soup = self.extractor.fetch_webpage(url) 
            if not full_text:
                msg = f"Contenido nulo en URL {i}, se descarta."
                print(msg)
                yield f"data: {msg}\n\n" # Notificamos el fallo
                log.append(msg)
                continue
                
            # 2. Extracción Estructurada
            structured_chunks = self.extractor._scrape_with_fallback(soup)
            if not structured_chunks:
                print(f"Fallo estructural en URL {i}, aplicando Plan C...")
                yield f"data: Fallo estructural en URL {i}...\n\n"
                structured_chunks = self._aggressive_text_fallback(soup, title)

            if not structured_chunks:
                log.append(f"Contenido insuficiente en {url}")
                continue

            # Consolidamos el contenido bruto
            url_consolidated_raw = "\n\n".join([
                f"--- SECCIÓN: {chunk_data['heading']} ---\n{chunk_data['content']}" 
                for chunk_data in structured_chunks
            ])
            
            all_media = [media for b in structured_chunks for media in b.get('media_elements', [])]
            
            # 3. FASE 3: Análisis de IA
            summary = ""
            if url_consolidated_raw:
                if req.use_ai:
                    # 📢 ENVIAR AL FRONTEND: "Analizando con IA..."
                    yield f"data: Analizando con IA URL {i}...\n\n"
                    
                    summary = self.ai_service.analizar_bloque_contenido(
                        url_consolidated_raw, 
                        all_media, 
                        analisis_contexto, 
                        title,
                        contexto_previo=consolidated_text 
                    )
                    consolidated_text += f"\n\n--- ANÁLISIS FUENTE: {url} ---\n{summary}"
                else:
                    summary = url_consolidated_raw
                    consolidated_text += f"\n\n{summary}"

            # Guardamos el resultado individual
            chunk_for_result = {
                'heading': title,
                'content': url_consolidated_raw,
                'media_elements': all_media,
                'ai_chunk_summary': summary
            }

            result = models.ScrapeResult(
                url=url,
                title=title,
                ai_titles=[],
                subtitles=[],
                text_content=summary,
                headers={"main_heading": [title], "count": [str(len(structured_chunks))]},
                ai_analysis=summary,
                title_suggestions=[], 
                final_structure="",
                article_blocks=[chunk_for_result], 
                status='OK' 
            )
            all_results.append(result)
            
            # 📢 ENVIAR AL FRONTEND: "URL X completada exitosamente"
            # Esto disparará el check verde en el frontend
            msg_exito = f"✔️ URL {i} completada exitosamente."
            print(msg_exito)
            yield f"data: {msg_exito}\n\n"

        # 4. FASE 4: Persistencia (Igual que tu código)
        valid_results = [r for r in all_results if r.status == 'OK'] 
        
        if valid_results and consolidated_text:
            try:
                # Tu lógica de guardado en DB...
                estimated_word_count_calculated = len(consolidated_text.split())
                all_article_blocks_combined = [b for res in all_results for b in (res.article_blocks or [])]
                
                datos_a_guardar_scraping = {
                    "consolidated_content": consolidated_text, 
                    'scrape_blocks_json': all_article_blocks_combined 
                }

                actualizar_o_crear_resultado_scraping(self.db, self.blog_id, datos_a_guardar_scraping) 
                
                blog_entity = self.db.query(Blog).filter(Blog.id == self.blog_id).first()
                estructura_actual = blog_entity.estructura_blog_json if blog_entity else {}
                
                actualizar_estructura_blog(
                    self.db, self.blog_id, 
                    estructura_data=estructura_actual, 
                    estimated_word_count=estimated_word_count_calculated
                )
            except Exception as e:
                print(f"ERROR DB: {str(e)} ")

        # 5. FASE 5: RESPUESTA FINAL
        final_response = models.ScrapeResponse(
            query=req.title_base if req.title_base else "",
            count=len(valid_results),
            results=all_results,
            log=log,
            consolidated_content=consolidated_text 
        )

        yield "event: final_data\n"
        yield f"data: {final_response.model_dump_json()}\n\n"


def execute_scraping(db: Session, blog_id: UUID, req: models.ScrapeRequest) -> Generator[str, None, None]:
    """
    Punto de entrada para la ejecución del scraping (Orquestador).
    Recibe la sesión DB y el ID del blog para la persistencia.
    """
    ai_service = AIService()
    extractor = ContentExtractor()
    # Pasa DB y blog_id al Orchestrator para que pueda guardar en la tabla 'scraping'
    orchestrator = AnalysisOrchestrator(ai_service, extractor, db, blog_id)
    return orchestrator.execute_scraping(req)

def analisis_final_ia(req: models.AIAnalysisRequest) -> Dict[str, Any]:
    """Punto de entrada para el análisis final de IA (Servicio de IA), incluyendo regeneración de secciones."""
    ai_service = AIService()
    return ai_service.analisis_final_ia(req)